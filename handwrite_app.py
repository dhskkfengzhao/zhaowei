import sys
import json
import docx
import fitz  # PyMuPDF
from pathlib import Path
from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6.QtGui import *
from PIL import Image, ImageFont, ImageQt, ImageDraw
from functools import partial
import time
import os
import threading
import zipfile
import shutil
import io

try:
    from handright import Template, handwrite
except ImportError:
    print("请先安装 handright: pip install handright pillow")
    sys.exit(1)

# 内置样式和默认设置
class StyleManager:
    """样式和默认设置管理器，直接将资源内嵌到代码中"""
    
    @staticmethod
    def get_default_style():
        """获取默认样式表内容"""
        return """
        /* 现代化的样式表 */
        QMainWindow {
            background-color: #f0f0f0;
        }

        QToolBar {
            border: none;
            background-color: #ffffff;
            spacing: 5px;
            padding: 5px;
        }

        QToolButton {
            border: none;
            border-radius: 4px;
            padding: 5px;
            background-color: transparent;
        }

        QToolButton:hover {
            background-color: #e0e0e0;
        }

        QDockWidget {
            border: 1px solid #cccccc;
            titlebar-close-icon: url(:/icons/close.png);
        }

        QGroupBox {
            border: 1px solid #cccccc;
            border-radius: 4px;
            margin-top: 1em;
            padding-top: 10px;
        }

        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 3px 0 3px;
        }

        QTextEdit {
            border: 1px solid #cccccc;
            border-radius: 4px;
            background-color: #ffffff;
        }

        QComboBox {
            border: 1px solid #cccccc;
            border-radius: 4px;
            padding: 5px;
            min-width: 6em;
        }

        QSpinBox {
            border: 1px solid #cccccc;
            border-radius: 4px;
            padding: 5px;
        }

        /* 深色主题样式将在代码中动态切换 */
        """
    
    @staticmethod
    def get_default_settings():
        """获取默认设置内容"""
        return {
            "font": "Bo Le Locust Tree Handwriting Pen Chinese Font-Simplified Chinese Fonts",
            "font_size": 40,
            "margins": {
                "上边距": 105,
                "下边距": 0,
                "左边距": 86,
                "右边距": 93,
                "字间距": 1,
                "行间距": 37
            },
            "distortions": {
                "字间距扰动": 2.0,
                "行间距扰动": 0.0,
                "字体大小扰动": 2.0,
                "横向偏移扰动": 2.0,
                "纵向偏移扰动": 2.0,
                "旋转角度扰动": 0.05
            }
        }
    
    @staticmethod
    def apply_default_style(app):
        """应用默认样式到应用程序"""
        app.setStyleSheet(StyleManager.get_default_style())

# 添加工作线程类，用于处理耗时的渲染操作
class WorkerThread(QThread):
    """工作线程类，用于后台处理耗时的渲染任务"""
    progressChanged = Signal(int)  # 进度更新信号
    resultReady = Signal(object)   # 结果就绪信号
    errorOccurred = Signal(str)    # 错误信号
    
    def __init__(self, text, template, parent=None):
        super().__init__(parent)
        self.text = text
        self.template = template
        self.is_cancelled = False
        
    def run(self):
        """执行渲染任务"""
        try:
            result = []
            # 检查必要的参数
            if not self.text:
                self.text = "预览文本示例"
            
            # 使用生成器获取页面，并更新进度
            pages = handwrite(self.text, self.template)
            
            # 设置进度监控器
            text_length = len(self.text)
            progress_step = max(1, text_length // 100)  # 每处理这么多字符更新一次进度
            chars_processed = 0
            
            for i, page in enumerate(pages):
                if self.is_cancelled:
                    break
                    
                result.append(page)
                
                # 更新进度 - 根据处理的字符数而不是页数
                chars_processed += progress_step
                progress = min(95, int(chars_processed / text_length * 100))
                self.progressChanged.emit(progress)
                
            if not self.is_cancelled:
                self.progressChanged.emit(100)
                self.resultReady.emit(result)
                
        except Exception as e:
            error_msg = str(e)
            # 提供更具体的错误提示
            if "font.size" in error_msg and "line_spacing" in error_msg:
                error_msg = "字体大小与行间距设置不合理。请确保行间距大于字体大小。"
            self.errorOccurred.emit(error_msg)
    
    def cancel(self):
        """取消操作"""
        self.is_cancelled = True

class PreviewManager:
    def __init__(self):
        self.last_update = 0
        self.throttle_delay = 200  # 节流延迟(毫秒)
        self.current_page = 0
        self.pages = []
        self.cached_images = {}  # 缓存已生成的图像
        self.last_settings = None  # 上次使用的设置
        self.last_text = ""  # 上次渲染的文本
        
    def should_update(self, text, settings):
        """判断是否需要更新预览"""
        # 检查文本和设置是否变化
        settings_changed = self.last_settings != settings
        text_changed = self.last_text != text
        
        current_time = time.time() * 1000
        if (current_time - self.last_update > self.throttle_delay) and (settings_changed or text_changed):
            self.last_update = current_time
            self.last_settings = settings.copy() if settings else None
            self.last_text = text
            return True
        return False
    
    def set_pages(self, pages):
        self.pages = list(pages)
        self.current_page = 0
    
    def get_current_page(self):
        if self.pages:
            return self.pages[self.current_page]
        return None
    
    def next_page(self):
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            return True
        return False
    
    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            return True
        return False
    
    def go_to_page(self, page_num):
        """跳转到指定页面"""
        if 0 <= page_num < len(self.pages):
            self.current_page = page_num
            return True
        return False
    
    def get_page_info(self):
        return f"第 {self.current_page + 1} 页,共 {len(self.pages)} 页"
    
    def get_total_pages(self):
        return len(self.pages)
    
    def clear_cache(self):
        """清除缓存"""
        self.cached_images.clear()

class PreviewWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.preview_manager = PreviewManager()
        self.zoom_level = 100  # 默认缩放级别为100%
        self.setup_ui()
        self.worker = None  # 保存当前的工作线程
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # 添加页面导航和缩放工具栏
        toolbar = QHBoxLayout()
        
        self.prev_btn = QPushButton("上一页")
        self.next_btn = QPushButton("下一页")
        self.page_label = QLabel()
        self.page_spin = QSpinBox()
        self.page_spin.setMinimum(1)
        self.page_spin.setValue(1)
        
        self.zoom_out_btn = QPushButton("缩小")
        self.zoom_in_btn = QPushButton("放大")
        self.zoom_label = QLabel("100%")
        self.fit_btn = QPushButton("适合窗口")
        
        self.prev_btn.clicked.connect(self.prev_page)
        self.next_btn.clicked.connect(self.next_page)
        self.page_spin.valueChanged.connect(self.go_to_page)
        self.zoom_out_btn.clicked.connect(lambda: self.set_zoom(self.zoom_level - 10))
        self.zoom_in_btn.clicked.connect(lambda: self.set_zoom(self.zoom_level + 10))
        self.fit_btn.clicked.connect(self.fit_to_window)
        
        toolbar.addWidget(self.prev_btn)
        toolbar.addWidget(self.page_spin)
        toolbar.addWidget(QLabel("/"))
        toolbar.addWidget(self.page_label)
        toolbar.addWidget(self.next_btn)
        toolbar.addStretch()
        toolbar.addWidget(self.zoom_out_btn)
        toolbar.addWidget(self.zoom_label)
        toolbar.addWidget(self.zoom_in_btn)
        toolbar.addWidget(self.fit_btn)
        
        layout.addLayout(toolbar)
        
        # 预览区域
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        
        self.preview_label = QLabel()
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setMinimumSize(600, 400)
        
        self.scroll_area.setWidget(self.preview_label)
        layout.addWidget(self.scroll_area)
        
        self.update_navigation()
    
    def update_preview(self, text, template):
        """更新预览内容"""
        # 检查是否有正在运行的线程，如果有则取消
        if self.worker and self.worker.isRunning():
            self.worker.cancel()
            self.worker.wait(100)  # 等待线程结束
        
        # 创建进度对话框
        progress_dialog = QProgressDialog("正在生成预览...", "取消", 0, 100, self)
        progress_dialog.setWindowModality(Qt.WindowModal)
        progress_dialog.setMinimumDuration(500)  # 仅当操作超过500ms时显示
        
        # 创建工作线程
        self.worker = WorkerThread(text, template)
        self.worker.progressChanged.connect(progress_dialog.setValue)
        self.worker.resultReady.connect(lambda images: self.handle_preview_result(images))
        self.worker.errorOccurred.connect(self.handle_preview_error)
        
        # 连接取消按钮
        progress_dialog.canceled.connect(self.worker.cancel)
        
        # 启动线程
        self.worker.start()
    
    def handle_preview_error(self, error_msg):
        """处理预览错误"""
        QMessageBox.warning(self, "预览错误", error_msg)
        
        # 如果没有页面可显示，显示一个错误图像
        if not self.preview_manager.pages:
            # 创建一个显示错误信息的图像
            error_image = Image.new('RGB', (800, 600), color='white')
            draw = ImageDraw.Draw(error_image)
            draw.text((50, 50), f"预览错误: {error_msg}", fill=(255, 0, 0))
            
            # 添加到预览管理器
            self.preview_manager.set_pages([error_image])
            self.show_current_page()
            self.update_navigation()
    
    def handle_preview_result(self, images):
        """处理预览结果"""
        if not images:
            self.handle_preview_error("未能生成预览图像，可能是参数设置不合理")
            return
            
        self.preview_manager.set_pages(images)
        self.page_spin.setMaximum(max(1, len(images)))
        self.show_current_page()
        self.update_navigation()
    
    def show_current_page(self):
        """显示当前页面"""
        current_image = self.preview_manager.get_current_page()
        if current_image:
            # 转换为Qt图像
            qt_image = ImageQt.ImageQt(current_image)
            pixmap = QPixmap.fromImage(qt_image)
            
            # 应用缩放
            self.apply_zoom(pixmap)
            
            # 更新页码
            self.page_spin.setValue(self.preview_manager.current_page + 1)
    
    def apply_zoom(self, pixmap):
        """应用缩放到图像"""
        if pixmap:
            # 计算缩放后的尺寸
            scaled_width = int(pixmap.width() * self.zoom_level / 100)
            scaled_height = int(pixmap.height() * self.zoom_level / 100)
            
            # 缩放图像
            scaled_pixmap = pixmap.scaled(
                scaled_width, 
                scaled_height,
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
            
            self.preview_label.setPixmap(scaled_pixmap)
            self.preview_label.resize(scaled_pixmap.size())
    
    def set_zoom(self, level):
        """设置缩放级别"""
        self.zoom_level = max(10, min(300, level))  # 限制在10%-300%之间
        self.zoom_label.setText(f"{self.zoom_level}%")
        current_image = self.preview_manager.get_current_page()
        if current_image:
            qt_image = ImageQt.ImageQt(current_image)
            pixmap = QPixmap.fromImage(qt_image)
            self.apply_zoom(pixmap)
    
    def fit_to_window(self):
        """适合窗口显示"""
        current_image = self.preview_manager.get_current_page()
        if current_image:
            qt_image = ImageQt.ImageQt(current_image)
            pixmap = QPixmap.fromImage(qt_image)
            
            # 计算合适的缩放比例
            view_width = self.scroll_area.viewport().width() - 20
            view_height = self.scroll_area.viewport().height() - 20
            
            image_width = pixmap.width()
            image_height = pixmap.height()
            
            width_ratio = view_width / image_width
            height_ratio = view_height / image_height
            
            # 使用较小的比例以确保完全适应
            ratio = min(width_ratio, height_ratio) * 100
            
            self.set_zoom(int(ratio))
    
    def update_navigation(self):
        """更新导航按钮状态"""
        has_pages = bool(self.preview_manager.pages)
        self.prev_btn.setEnabled(has_pages and self.preview_manager.current_page > 0)
        self.next_btn.setEnabled(has_pages and self.preview_manager.current_page < len(self.preview_manager.pages) - 1)
        self.page_label.setText(str(len(self.preview_manager.pages)) if has_pages else "0")
        self.page_spin.setEnabled(has_pages)
    
    def prev_page(self):
        """显示上一页"""
        if self.preview_manager.prev_page():
            self.show_current_page()
            self.update_navigation()
    
    def next_page(self):
        """显示下一页"""
        if self.preview_manager.next_page():
            self.show_current_page()
            self.update_navigation()
    
    def go_to_page(self, page_num):
        """跳转到指定页面"""
        if page_num != (self.preview_manager.current_page + 1):
            if self.preview_manager.go_to_page(page_num - 1):
                self.show_current_page()
                self.update_navigation()

class SettingsPanel(QWidget):
    settingsChanged = Signal()
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # 创建分组
        self.create_font_group()
        self.create_margin_group()
        self.create_distortion_group()
        self.create_preset_group()
        
        layout.addStretch()
    
    def create_font_group(self):
        group = QGroupBox("字体设置")
        layout = QFormLayout(group)
        
        self.font_combo = QComboBox()
        self.update_font_list()
        
        self.font_size = QSpinBox()
        self.font_size.setRange(1, 200)
        self.font_size.setValue(100)
        
        layout.addRow("字体:", self.font_combo)
        layout.addRow("字号:", self.font_size)
        
        self.layout().addWidget(group)
    
    def create_margin_group(self):
        group = QGroupBox("边距设置")
        layout = QGridLayout(group)
        
        self.margin_inputs = {}
        margins = [
            ("上边距", 140),  # 保持默认上边距
            ("下边距", 70),   # 添加下边距设置
            ("左边距", 18),   # 调整为更合理的左边距
            ("右边距", 100),  # 添加右边距设置
            ("字间距", 5),    # 保持默认字间距
            ("行间距", 100)   # 增加行间距，确保大于字体大小
        ]
        
        for i, (name, default) in enumerate(margins):
            spinbox = QSpinBox()
            spinbox.setRange(0, 500)
            
            # 对于行间距，设置最小值为60，避免太小
            if name == "行间距":
                spinbox.setMinimum(60)
                spinbox.setToolTip("行间距必须大于字体大小，至少为60")
            
            spinbox.setValue(int(default))
            spinbox.valueChanged.connect(self.on_settings_changed)
            self.margin_inputs[name] = spinbox
            layout.addWidget(QLabel(name), i//2, i%2*2)
            layout.addWidget(spinbox, i//2, i%2*2+1)
            
        self.layout().addWidget(group)
    
    def create_distortion_group(self):
        group = QGroupBox("扰动设置")
        layout = QFormLayout(group)
        
        self.distortion_inputs = {}
        distortions = [
            ("字间距扰动", 2.0),
            ("行间距扰动", 1.0),      # 增加行间距扰动，使其更自然
            ("字体大小随机扰动", 2.0),
            ("字体大小", 40),         # 减小默认字体大小，确保小于行间距
            ("笔画横向偏移随机扰动", 2.0),  # 减小横向扰动
            ("笔画纵向偏移随机扰动", 2.0),  # 减小纵向扰动
            ("笔画旋转偏移随机扰动", 0.05)  # 保持小的旋转扰动
        ]
        
        for name, default in distortions:
            spinbox = QDoubleSpinBox()
            spinbox.setRange(0, 10)
            spinbox.setSingleStep(0.1)
            spinbox.setValue(default)
            spinbox.valueChanged.connect(self.on_settings_changed)
            self.distortion_inputs[name] = spinbox
            layout.addRow(name + ":", spinbox)
            
        self.layout().addWidget(group)
    
    def create_preset_group(self):
        group = QGroupBox("预设管理")
        layout = QHBoxLayout(group)
        
        self.preset_combo = QComboBox()
        self.preset_combo.addItems(["默认"])
        
        save_btn = QPushButton("保存预设")
        save_btn.clicked.connect(self.save_preset)
        
        delete_btn = QPushButton("删除预设")
        delete_btn.clicked.connect(self.delete_preset)
        
        layout.addWidget(self.preset_combo)
        layout.addWidget(save_btn)
        layout.addWidget(delete_btn)
        
        self.layout().addWidget(group)
    
    def update_font_list(self):
        """更新字体列表"""
        self.font_combo.clear()
        fonts_dir = Path("fonts")
        if fonts_dir.exists():
            for font in fonts_dir.glob("*.ttf"):
                self.font_combo.addItem(font.stem)
        if self.font_combo.count() == 0:
            self.font_combo.addItem("默认字体")
    
    def get_current_settings(self):
        """获取当前所有设置"""
        return {
            "font": self.font_combo.currentText(),
            "font_size": self.font_size.value(),
            "margins": {k: v.value() for k, v in self.margin_inputs.items()},
            "distortions": {k: v.value() for k, v in self.distortion_inputs.items()}
        }
    
    def on_settings_changed(self):
        """设置变更时发出信号"""
        self.settingsChanged.emit()
    
    def save_preset(self):
        """保存当前设置为预设"""
        name, ok = QInputDialog.getText(self, "保存预设", "请输入预设名称:")
        if ok and name:
            settings = self.get_current_settings()
            # 保存到文件
            presets_file = Path("presets.json")
            presets = {}
            if presets_file.exists():
                with open(presets_file, 'r', encoding='utf-8') as f:
                    presets = json.load(f)
            presets[name] = settings
            with open(presets_file, 'w', encoding='utf-8') as f:
                json.dump(presets, f, ensure_ascii=False, indent=2)
            
            if self.preset_combo.findText(name) == -1:
                self.preset_combo.addItem(name)
    
    def delete_preset(self):
        """删除当前选中的预设"""
        current = self.preset_combo.currentText()
        if current == "默认":
            return
            
        reply = QMessageBox.question(self, "确认删除", 
                                   f"确定要删除预设 {current} 吗？",
                                   QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            presets_file = Path("presets.json")
            if presets_file.exists():
                with open(presets_file, 'r', encoding='utf-8') as f:
                    presets = json.load(f)
                if current in presets:
                    del presets[current]
                    with open(presets_file, 'w', encoding='utf-8') as f:
                        json.dump(presets, f, ensure_ascii=False, indent=2)
            
            self.preset_combo.removeItem(self.preset_combo.currentIndex())

class SettingsManager:
    def __init__(self):
        self.settings_file = Path("settings.json")
        self.presets_file = Path("presets.json")
        
    def load_settings(self):
        """加载设置，如果找不到设置文件则使用默认设置"""
        if self.settings_file.exists():
            try:
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"加载设置失败: {e}")
                return StyleManager.get_default_settings()
        else:
            # 使用内置的默认设置
            return StyleManager.get_default_settings()
        
    def save_settings(self, settings):
        """保存当前设置"""
        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存设置失败: {e}")
    
    def load_presets(self):
        """加载预设"""
        if self.presets_file.exists():
            try:
                with open(self.presets_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"加载预设失败: {e}")
        return {}
    
    def save_preset(self, name, settings):
        """保存预设"""
        presets = self.load_presets()
        presets[name] = settings
        try:
            with open(self.presets_file, 'w', encoding='utf-8') as f:
                json.dump(presets, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"保存预设失败: {e}")
            return False
    
    def get_default_settings(self):
        """获取默认设置"""
        return StyleManager.get_default_settings()

class ThemeManager:
    """主题管理器，用于切换应用主题"""
    
    def __init__(self, app):
        self.app = app
        self.current_theme = "light"  # 默认使用浅色主题
        
    def set_theme(self, theme):
        """设置应用主题
        
        Args:
            theme: 主题名称，可选值为 "light" 或 "dark"
        """
        self.current_theme = theme
        
        if theme == "light":
            self.app.setStyleSheet(self._get_light_theme())
        else:
            self.app.setStyleSheet(self._get_dark_theme())
    
    def toggle_theme(self):
        """切换主题"""
        if self.current_theme == "light":
            self.set_theme("dark")
        else:
            self.set_theme("light")
    
    def _get_light_theme(self):
        """获取浅色主题样式表"""
        return StyleManager.get_default_style()
    
    def _get_dark_theme(self):
        """获取深色主题样式表"""
        return """
        /* 深色主题样式表 */
        QMainWindow, QDialog {
            background-color: #2d2d2d;
        }
        
        QWidget {
            color: #e0e0e0;
        }
        
        QToolBar {
            border: none;
            background-color: #383838;
            spacing: 5px;
            padding: 5px;
        }
        
        QPushButton {
            background-color: #505050;
            border: 1px solid #606060;
            border-radius: 4px;
            padding: 5px 10px;
            color: #e0e0e0;
        }
        
        QPushButton:hover {
            background-color: #606060;
        }
        
        QPushButton:pressed {
            background-color: #707070;
        }
        
        QToolButton {
            border: none;
            border-radius: 4px;
            padding: 5px;
            background-color: transparent;
        }
        
        QToolButton:hover {
            background-color: #505050;
        }
        
        QDockWidget {
            border: 1px solid #505050;
        }
        
        QGroupBox {
            border: 1px solid #505050;
            border-radius: 4px;
            margin-top: 1em;
            padding-top: 10px;
            background-color: #383838;
        }
        
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 3px 0 3px;
        }
        
        QTextEdit, QPlainTextEdit {
            border: 1px solid #505050;
            border-radius: 4px;
            background-color: #383838;
            selection-background-color: #3d6185;
            color: #e0e0e0;
        }
        
        QComboBox, QSpinBox, QDoubleSpinBox {
            border: 1px solid #505050;
            border-radius: 4px;
            padding: 5px;
            background-color: #383838;
            color: #e0e0e0;
        }
        
        QComboBox:focus, QSpinBox:focus, QDoubleSpinBox:focus {
            border-color: #5c88c5;
        }
        
        QScrollBar:vertical {
            border: none;
            background: #383838;
            width: 10px;
            margin: 0px;
        }
        
        QScrollBar::handle:vertical {
            background: #606060;
            min-height: 20px;
            border-radius: 5px;
        }
        
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
            height: 0px;
        }
        
        QScrollBar:horizontal {
            border: none;
            background: #383838;
            height: 10px;
            margin: 0px;
        }
        
        QScrollBar::handle:horizontal {
            background: #606060;
            min-width: 20px;
            border-radius: 5px;
        }
        
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
            width: 0px;
        }
        
        QStatusBar {
            background-color: #383838;
            color: #a0a0a0;
        }
        """

class FileManager:
    @staticmethod
    def import_file(file_path):
        """导入各种格式的文件
        
        Args:
            file_path: 文件路径对象
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 导入失败时抛出异常
        """
        ext = file_path.suffix.lower()
        text = ""
        
        try:
            if ext == '.txt':
                # 尝试多种编码
                encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ascii']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        break  # 如果成功读取，跳出循环
                    except UnicodeDecodeError:
                        continue  # 尝试下一种编码
                
                if not text:  # 如果所有编码都失败
                    raise ValueError(f"无法解码文件 {file_path}，请检查文件编码")
                    
            elif ext == '.docx':
                doc = docx.Document(file_path)
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # 只添加非空段落
                        paragraphs.append(paragraph.text)
                text = '\n'.join(paragraphs)
                
            elif ext == '.pdf':
                doc = fitz.open(file_path)
                pages_text = []
                for page in doc:
                    page_text = page.get_text()
                    if page_text.strip():  # 只添加非空页面
                        pages_text.append(page_text)
                text = '\n'.join(pages_text)
                
            elif ext in ['.doc', '.xls', '.ppt']:
                raise ValueError(f"不支持旧版Office格式 ({ext})，请转换为新格式后再试")
                
            else:
                raise ValueError(f"不支持的文件格式: {ext}")
                
            return text
        except Exception as e:
            raise Exception(f"导入文件失败: {str(e)}")
    
    @staticmethod
    def export_images(images, base_path, format='png', dpi=300):
        """导出图片
        
        Args:
            images: 图片列表
            base_path: 基础路径
            format: 导出格式 (png, jpg, pdf)
            dpi: 分辨率 (仅用于PDF)
            
        Returns:
            list: 导出的文件路径列表
            
        Raises:
            Exception: 导出失败时抛出异常
        """
        results = []
        try:
            if format == 'pdf':
                # 如果是PDF格式，将所有图片合并为一个PDF文件
                pdf_path = f"{base_path}.pdf"
                
                # 确保第一张图片存在
                if images:
                    # 保存为PDF
                    images[0].save(
                        pdf_path, 
                        "PDF", 
                        resolution=float(dpi), 
                        save_all=True, 
                        append_images=images[1:] if len(images) > 1 else []
                    )
                    results.append(pdf_path)
                
            elif format.lower() in ['png', 'jpg', 'jpeg']:
                # 导出为单独的图片文件
                ext = 'jpg' if format.lower() in ['jpg', 'jpeg'] else 'png'
                img_format = 'JPEG' if ext == 'jpg' else 'PNG'
                
                for i, image in enumerate(images):
                    page_path = f"{base_path}_第{i+1}页.{ext}"
                    image.save(page_path, format=img_format)
                    results.append(page_path)
                    
            elif format.lower() == 'docx':
                # 将图片导出为Word文档
                docx_path = f"{base_path}.docx"
                doc = docx.Document()
                
                for i, image in enumerate(images):
                    # 保存为临时文件
                    temp_path = f"{base_path}_temp_{i}.png"
                    image.save(temp_path, 'PNG')
                    
                    # 添加到Word文档
                    doc.add_picture(temp_path, width=docx.shared.Inches(6))
                    
                    # 如果不是最后一页，添加分页符
                    if i < len(images) - 1:
                        doc.add_page_break()
                    
                    # 删除临时文件
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                
                # 保存Word文档
                doc.save(docx_path)
                results.append(docx_path)
                
            else:
                raise ValueError(f"不支持的导出格式: {format}")
                
            return results
        except Exception as e:
            raise Exception(f"导出图片失败: {str(e)}")
    
    @staticmethod
    def get_supported_import_formats():
        """获取支持的导入格式"""
        return "文本文件 (*.txt);;Word文档 (*.docx);;PDF文件 (*.pdf);;所有文件 (*.*)"
    
    @staticmethod
    def get_supported_export_formats():
        """获取支持的导出格式"""
        return "PDF文件 (*.pdf);;PNG图片 (*.png);;JPEG图片 (*.jpg);;Word文档 (*.docx)"

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("手写模拟器 v1.0")
        self.settings_manager = SettingsManager()
        self.setup_ui()
        self.setup_connections()
        self.load_settings()
        
        # 创建状态栏
        self.statusBar().showMessage("就绪")
        
        # 创建工作线程对象
        self.worker_thread = None
        
        # 主题管理器
        self.theme_manager = ThemeManager(QApplication.instance())
        
    def setup_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 创建工具栏
        toolbar = self.addToolBar("主工具栏")
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(24, 24))
        
        # 文件操作
        import_action = QAction("导入文件", self)
        import_action.setStatusTip("导入文本文件")
        import_action.triggered.connect(self.import_text)
        toolbar.addAction(import_action)
        
        export_action = QAction("导出", self)
        export_action.setStatusTip("导出手写效果")
        export_action.triggered.connect(self.export_image)
        toolbar.addAction(export_action)
        
        toolbar.addSeparator()
        
        # 设置操作
        save_settings_action = QAction("保存设置", self)
        save_settings_action.setStatusTip("保存当前设置")
        save_settings_action.triggered.connect(self.save_settings)
        toolbar.addAction(save_settings_action)
        
        load_preset_action = QAction("加载预设", self)
        load_preset_action.setStatusTip("加载预设配置")
        load_preset_action.triggered.connect(self.load_preset_dialog)
        toolbar.addAction(load_preset_action)
        
        save_preset_action = QAction("保存预设", self)
        save_preset_action.setStatusTip("保存当前配置为预设")
        save_preset_action.triggered.connect(self.save_preset_dialog)
        toolbar.addAction(save_preset_action)
        
        toolbar.addSeparator()
        
        # 主题切换
        theme_action = QAction("切换主题", self)
        theme_action.setStatusTip("切换明暗主题")
        theme_action.triggered.connect(self.toggle_theme)
        toolbar.addAction(theme_action)
        
        toolbar.addSeparator()
        
        # 帮助
        help_action = QAction("帮助", self)
        help_action.setStatusTip("查看帮助信息")
        help_action.triggered.connect(self.show_help)
        toolbar.addAction(help_action)
        
        about_action = QAction("关于", self)
        about_action.setStatusTip("关于手写模拟器")
        about_action.triggered.connect(self.show_about)
        toolbar.addAction(about_action)
        
        # 主要内容区域
        content_layout = QHBoxLayout()
        
        # 左侧设置面板
        settings_panel = QWidget()
        settings_layout = QVBoxLayout(settings_panel)
        
        # 基本设置组
        basic_group = QGroupBox("基本设置")
        basic_layout = QFormLayout(basic_group)
        
        # 字体和背景选择
        self.font_combo = QComboBox()
        self.bg_combo = QComboBox()
        self.update_font_list()
        self.update_background_list()
        basic_layout.addRow("字体:", self.font_combo)
        basic_layout.addRow("背景:", self.bg_combo)
        settings_layout.addWidget(basic_group)
        
        # 边距设置组
        margin_group = QGroupBox("边距设置")
        margin_layout = QFormLayout(margin_group)
        
        # 边距设置
        self.margin_inputs = {}
        margins = [
            ("上边距", 140), 
            ("下边距", 70),  
            ("左边距", 18),  
            ("右边距", 100), 
            ("字间距", 5),   
            ("行间距", 100)  
        ]
        
        for name, default in margins:
            spinbox = QSpinBox()
            spinbox.setRange(0, 500)
            
            # 对于行间距，设置最小值为60，避免太小
            if name == "行间距":
                spinbox.setMinimum(60)
                spinbox.setToolTip("行间距必须大于字体大小，至少为60")
            
            spinbox.setValue(default)
            spinbox.valueChanged.connect(self.delayed_preview_update)
            self.margin_inputs[name] = spinbox
            margin_layout.addRow(f"{name}:", spinbox)
        
        settings_layout.addWidget(margin_group)
        
        # 扰动设置组
        distortion_group = QGroupBox("扰动设置")
        distortion_layout = QFormLayout(distortion_group)
        
        # 扰动设置
        self.distortion_inputs = {}
        distortions = [
            ("字间距扰动", 2.0),
            ("行间距扰动", 1.0),     
            ("字体大小随机扰动", 2.0),
            ("字体大小", 40),        
            ("笔画横向偏移随机扰动", 2.0), 
            ("笔画纵向偏移随机扰动", 2.0), 
            ("笔画旋转偏移随机扰动", 0.05) 
        ]
        
        for name, default in distortions:
            if "扰动" in name:
                spinbox = QDoubleSpinBox()
                spinbox.setRange(0, 10)
                spinbox.setSingleStep(0.1)
            else:
                spinbox = QSpinBox()
                spinbox.setRange(1, 500)
                
                # 添加工具提示说明
                if name == "字体大小":
                    spinbox.setToolTip("字体大小增加时，行间距也会自动增加")
            
            spinbox.setValue(default)
            spinbox.valueChanged.connect(self.delayed_preview_update)
            self.distortion_inputs[name] = spinbox
            distortion_layout.addRow(f"{name}:", spinbox)
        
        settings_layout.addWidget(distortion_group)
        
        # 批量处理设置
        batch_group = QGroupBox("批量处理")
        batch_layout = QVBoxLayout(batch_group)
        
        # 批量导入按钮
        batch_import_btn = QPushButton("批量导入文件")
        batch_import_btn.clicked.connect(self.batch_import)
        batch_layout.addWidget(batch_import_btn)
        
        # 批量导出按钮
        batch_export_btn = QPushButton("批量导出")
        batch_export_btn.clicked.connect(self.batch_export)
        batch_layout.addWidget(batch_export_btn)
        
        settings_layout.addWidget(batch_group)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        settings_layout.addWidget(self.progress_bar)
        
        # 添加伸缩项，确保设置面板紧凑
        settings_layout.addStretch()
        
        # 设置左侧面板的固定宽度
        settings_panel.setFixedWidth(300)
        content_layout.addWidget(settings_panel)
        
        # 右侧预览区域
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # 视图切换按钮组
        view_group = QHBoxLayout()
        self.radio_text = QRadioButton("文本")
        self.radio_preview = QRadioButton("预览")
        self.radio_text.setChecked(True)
        view_group.addWidget(self.radio_text)
        view_group.addWidget(self.radio_preview)
        view_group.addStretch()
        
        # 添加文本操作按钮
        clear_btn = QPushButton("清空")
        clear_btn.clicked.connect(self.clear_text)
        view_group.addWidget(clear_btn)
        
        # 添加快速预览按钮
        preview_btn = QPushButton("预览")
        preview_btn.clicked.connect(self.force_preview_update)
        view_group.addWidget(preview_btn)
        
        right_layout.addLayout(view_group)
        
        # 堆叠窗口
        self.stack = QStackedWidget()
        
        # 文本编辑区
        self.text_edit = QTextEdit()
        self.text_edit.setAcceptRichText(False)  # 只接受纯文本
        self.text_edit.setLineWrapMode(QTextEdit.WidgetWidth)
        self.stack.addWidget(self.text_edit)
        
        # 预览区
        self.preview = PreviewWidget()
        self.stack.addWidget(self.preview)
        
        right_layout.addWidget(self.stack)
        content_layout.addWidget(right_panel)
        
        main_layout.addLayout(content_layout)
        
        # 设置窗口大小和标题
        self.resize(1000, 800)
        
    def setup_connections(self):
        """设置信号连接"""
        # 文本变化时更新预览
        self.text_edit.textChanged.connect(self.delayed_preview_update)
        
        # 设置变化时更新预览
        for spinbox in self.margin_inputs.values():
            spinbox.valueChanged.connect(self.delayed_preview_update)
        for spinbox in self.distortion_inputs.values():
            spinbox.valueChanged.connect(self.delayed_preview_update)
        
        # 字体大小变化时，确保行间距合理
        if "字体大小" in self.distortion_inputs:
            self.distortion_inputs["字体大小"].valueChanged.connect(self.adjust_line_spacing)
        
        # 字体和背景变化时更新预览
        self.font_combo.currentIndexChanged.connect(self.delayed_preview_update)
        self.bg_combo.currentIndexChanged.connect(self.delayed_preview_update)
        
        # 预览/文本切换
        self.radio_text.toggled.connect(self.on_view_changed)
        self.radio_preview.toggled.connect(self.on_view_changed)
        
        # 创建定时器用于延迟更新
        self.preview_timer = QTimer()
        self.preview_timer.setSingleShot(True)
        self.preview_timer.timeout.connect(self.update_preview)
    
    def adjust_line_spacing(self, font_size):
        """当字体大小变化时，确保行间距足够大"""
        if "行间距" in self.margin_inputs:
            current_line_spacing = self.margin_inputs["行间距"].value()
            min_line_spacing = int(font_size * 1.5)  # 行间距至少为字体大小的1.5倍
            
            if current_line_spacing < min_line_spacing:
                self.margin_inputs["行间距"].setValue(min_line_spacing)
                self.statusBar().showMessage(f"已自动调整行间距为 {min_line_spacing} (必须大于字体大小)")
    
    def delayed_preview_update(self):
        """延迟更新预览，避免频繁更新"""
        # 如果不在预览模式，不更新
        if not self.radio_preview.isChecked():
            return
            
        # 设置状态栏提示
        self.statusBar().showMessage("预览内容已更改，等待更新...")
        
        # 启动延迟定时器，延迟500ms
        self.preview_timer.start(500)
    
    def force_preview_update(self):
        """强制更新预览"""
        # 切换到预览模式
        self.radio_preview.setChecked(True)
        # 立即更新预览
        self.update_preview()
    
    def update_preview(self):
        """更新预览图像"""
        try:
            if not self.radio_preview.isChecked():
                self.stack.setCurrentIndex(0)  # 切换到文本编辑视图
                return
                
            # 确保视图已切换到预览
            self.stack.setCurrentIndex(1)
                
            text = self.text_edit.toPlainText()
            if not text:
                text = "预览文本示例"
            
            # 获取当前设置
            settings = self.get_current_settings()
            
            # 创建模板
            template = self.create_template(settings)
            
            # 生成预览图像
            self.preview.update_preview(text[:min(500, len(text))], template)
            
            # 更新状态栏
            self.statusBar().showMessage("预览已更新")
            
        except Exception as e:
            self.statusBar().showMessage(f"预览错误: {str(e)}")
            QMessageBox.warning(self, "预览错误", str(e))
    
    def import_text(self):
        """导入文件"""
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "导入文件",
            "",
            FileManager.get_supported_import_formats()
        )
        
        if file_name:
            try:
                # 显示状态栏信息
                self.statusBar().showMessage(f"正在导入文件: {Path(file_name).name}...")
                
                # 导入文件
                text = FileManager.import_file(Path(file_name))
                
                # 如果已经有文本，询问是否替换或追加
                if self.text_edit.toPlainText():
                    reply = QMessageBox.question(
                        self,
                        "导入方式",
                        "已有文本，您希望如何处理？",
                        QMessageBox.Cancel | QMessageBox.Discard | QMessageBox.Save,
                        QMessageBox.Save
                    )
                    
                    if reply == QMessageBox.Save:  # 追加
                        self.text_edit.append("\n\n" + text)
                    elif reply == QMessageBox.Discard:  # 替换
                        self.text_edit.setPlainText(text)
                    # Cancel 则不做任何操作
                else:
                    self.text_edit.setPlainText(text)
                
                self.statusBar().showMessage(f"成功导入文件: {Path(file_name).name}")
                
            except Exception as e:
                self.statusBar().showMessage("导入失败")
                QMessageBox.critical(self, "导入错误", str(e))
    
    def export_image(self):
        """导出手写图片"""
        text = self.text_edit.toPlainText()
        if not text:
            QMessageBox.warning(self, "警告", "请先输入要转换的文本")
            return
        
        try:
            # 选择导出格式和位置
            file_dialog = QFileDialog(self)
            file_dialog.setAcceptMode(QFileDialog.AcceptSave)
            file_dialog.setWindowTitle("导出图片")
            file_dialog.setNameFilter(FileManager.get_supported_export_formats())
            
            if not file_dialog.exec():
                return
                
            file_name = file_dialog.selectedFiles()[0]
            selected_filter = file_dialog.selectedNameFilter()
            
            # 获取选择的格式
            if "PDF" in selected_filter:
                format = 'pdf'
            elif "PNG" in selected_filter:
                format = 'png'
            elif "JPEG" in selected_filter:
                format = 'jpeg'
            elif "DOCX" in selected_filter:
                format = 'docx'
            else:
                format = 'pdf'  # 默认PDF
            
            # 获取当前设置并检查设置是否合理
            settings = self.get_current_settings()
            
            # 检查字体大小和行间距
            font_size = settings['font_size']
            line_spacing = settings['margins']['行间距']
            
            if line_spacing <= font_size:
                # 自动调整行间距
                old_line_spacing = line_spacing
                line_spacing = int(font_size * 1.5)
                settings['margins']['行间距'] = line_spacing
                
                # 更新UI中的行间距值
                if '行间距' in self.margin_inputs:
                    self.margin_inputs['行间距'].setValue(line_spacing)
                
                # 显示警告
                QMessageBox.warning(
                    self,
                    "参数已调整",
                    f"行间距值({old_line_spacing})小于字体大小({font_size})，已自动调整为{line_spacing}。"
                )
            
            # 创建进度对话框
            progress_dialog = QProgressDialog("正在生成手写图片...", "取消", 0, 100, self)
            progress_dialog.setWindowModality(Qt.WindowModal)
            progress_dialog.setMinimumDuration(500)  # 仅当操作超过500ms时显示
            
            # 显示进度条
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            
            # 创建模板
            template = self.create_template(settings)
            
            # 创建工作线程
            self.worker_thread = WorkerThread(text, template)
            
            # 连接信号
            self.worker_thread.progressChanged.connect(self.progress_bar.setValue)
            self.worker_thread.progressChanged.connect(progress_dialog.setValue)
            
            # 定义结果处理函数
            def handle_result(images):
                try:
                    if not images:
                        QMessageBox.critical(self, "导出错误", "未能生成任何图像，请检查参数设置")
                        self.progress_bar.setVisible(False)
                        return
                        
                    # 导出图片
                    base_path = os.path.splitext(file_name)[0]
                    exported_files = FileManager.export_images(images, base_path, format)
                    
                    # 隐藏进度条
                    self.progress_bar.setVisible(False)
                    
                    # 显示成功消息
                    QMessageBox.information(
                        self,
                        "导出成功",
                        f"已导出 {len(exported_files)} 个文件:\n" + "\n".join(exported_files)
                    )
                except Exception as e:
                    self.progress_bar.setVisible(False)
                    QMessageBox.critical(self, "导出错误", str(e))
            
            # 定义错误处理函数
            def handle_error(error_msg):
                self.progress_bar.setVisible(False)
                
                # 提供更具体的错误提示和解决方案
                if "font.size" in error_msg and "line_spacing" in error_msg:
                    error_msg = "字体大小与行间距设置不合理。请确保行间距大于字体大小，建议行间距至少为字体大小的1.5倍。"
                    
                QMessageBox.critical(self, "导出错误", error_msg)
            
            # 连接结果和错误信号
            self.worker_thread.resultReady.connect(handle_result)
            self.worker_thread.errorOccurred.connect(handle_error)
            
            # 连接取消按钮
            progress_dialog.canceled.connect(self.worker_thread.cancel)
            
            # 启动线程
            self.worker_thread.start()
            
        except Exception as e:
            self.progress_bar.setVisible(False)
            QMessageBox.critical(self, "导出错误", str(e))
    
    def create_template(self, settings=None):
        """根据设置创建模板"""
        if settings is None:
            settings = self.get_current_settings()
        
        # 获取背景
        background = self.get_current_background()
        
        # 获取字体
        font_name = settings['font']
        font_path = Path("fonts") / f"{font_name}.ttf"
        if not font_path.exists():
            # 尝试查找完整文件名
            for font_file in Path("fonts").glob("*.ttf"):
                if font_file.stem == font_name:
                    font_path = font_file
                    break
            else:
                font_path = "simsun.ttc"  # 使用系统默认字体
        
        try:
            # 确保行间距始终大于字体大小，防止Handright报错
            font_size = settings['font_size']
            line_spacing = settings['margins']['行间距']
            
            # 行间距必须大于字体大小，至少为字体大小的1.5倍
            if line_spacing <= font_size:
                line_spacing = int(font_size * 1.5)
                # 更新UI中的行间距值
                if '行间距' in self.margin_inputs:
                    self.margin_inputs['行间距'].setValue(line_spacing)
                    self.statusBar().showMessage(f"已自动调整行间距为 {line_spacing} (必须大于字体大小)")
            
            return Template(
                background=background,
                font=ImageFont.truetype(str(font_path), size=font_size),
                line_spacing=line_spacing,
                word_spacing=settings['margins']['字间距'],
                left_margin=settings['margins']['左边距'],
                top_margin=settings['margins']['上边距'],
                right_margin=settings['margins']['右边距'],
                bottom_margin=settings['margins']['下边距'],
                word_spacing_sigma=settings['distortions']['字间距扰动'],
                line_spacing_sigma=settings['distortions']['行间距扰动'],
                font_size_sigma=settings['distortions']['字体大小扰动'],
                perturb_x_sigma=settings['distortions']['横向偏移扰动'],
                perturb_y_sigma=settings['distortions']['纵向偏移扰动'],
                perturb_theta_sigma=settings['distortions']['旋转角度扰动']
            )
        except Exception as e:
            raise Exception(f"创建模板失败: {str(e)}")
        
    def load_settings(self):
        """加载上次的设置"""
        settings = self.settings_manager.load_settings()
        if settings:
            try:
                # 恢复各项设置
                if 'font' in settings:
                    index = self.font_combo.findText(settings['font'])
                    if index >= 0:
                        self.font_combo.setCurrentIndex(index)
                
                if 'margins' in settings:
                    for name, value in settings['margins'].items():
                        if name in self.margin_inputs:
                            self.margin_inputs[name].setValue(value)
                
                if 'distortions' in settings:
                    for name, value in settings['distortions'].items():
                        mapped_name = {
                            "字间距扰动": "字间距扰动",
                            "行间距扰动": "行间距扰动",
                            "字体大小扰动": "字体大小随机扰动",
                            "横向偏移扰动": "笔画横向偏移随机扰动",
                            "纵向偏移扰动": "笔画纵向偏移随机扰动",
                            "旋转角度扰动": "笔画旋转偏移随机扰动"
                        }.get(name, name)
                        
                        if mapped_name in self.distortion_inputs:
                            self.distortion_inputs[mapped_name].setValue(value)
                            
                self.statusBar().showMessage("已加载上次的设置")
            except Exception as e:
                self.statusBar().showMessage(f"加载设置失败: {str(e)}")
    
    def save_settings(self):
        """保存当前设置"""
        try:
            settings = self.get_current_settings()
            self.settings_manager.save_settings(settings)
            self.statusBar().showMessage("设置已保存")
        except Exception as e:
            self.statusBar().showMessage(f"保存设置失败: {str(e)}")
    
    def closeEvent(self, event):
        """窗口关闭时保存设置"""
        self.save_settings()
        
        # 如果有正在运行的线程，停止它
        if self.worker_thread and self.worker_thread.isRunning():
            self.worker_thread.cancel()
            self.worker_thread.wait(1000)  # 等待最多1秒让线程结束
            
        super().closeEvent(event)

    def update_font_list(self):
        """更新字体列表"""
        self.font_combo.clear()
        fonts_dir = Path("fonts")
        if fonts_dir.exists():
            for font in fonts_dir.glob("*.ttf"):
                self.font_combo.addItem(font.stem)
        if self.font_combo.count() == 0:
            self.font_combo.addItem("默认字体")
    
    def update_background_list(self):
        """更新背景列表"""
        self.bg_combo.clear()
        backgrounds_dir = Path("backgrounds")
        self.bg_combo.addItem("默认白色背景", None)
        
        if backgrounds_dir.exists():
            for ext in ['.png', '.jpg', '.jpeg']:
                for f in backgrounds_dir.glob(f"*{ext}"):
                    self.bg_combo.addItem(f.stem, str(f))
    
    def get_current_background(self):
        """获取当前选中的背景"""
        bg_path = self.bg_combo.currentData()
        if bg_path and Path(bg_path).exists():
            return Image.open(bg_path).convert('RGB')
        return Image.new('RGB', (1000, 1000), color='white')
    
    def get_current_settings(self):
        """获取当前所有设置"""
        return {
            "font": self.font_combo.currentText(),
            "font_size": self.distortion_inputs["字体大小"].value(),
            "margins": {
                "上边距": self.margin_inputs["上边距"].value(),
                "下边距": self.margin_inputs["下边距"].value(),
                "左边距": self.margin_inputs["左边距"].value(),
                "右边距": self.margin_inputs["右边距"].value(),
                "字间距": self.margin_inputs["字间距"].value(),
                "行间距": self.margin_inputs["行间距"].value()
            },
            "distortions": {
                "字间距扰动": self.distortion_inputs["字间距扰动"].value(),
                "行间距扰动": self.distortion_inputs["行间距扰动"].value(),
                "字体大小扰动": self.distortion_inputs["字体大小随机扰动"].value(),
                "横向偏移扰动": self.distortion_inputs["笔画横向偏移随机扰动"].value(),
                "纵向偏移扰动": self.distortion_inputs["笔画纵向偏移随机扰动"].value(),
                "旋转角度扰动": self.distortion_inputs["笔画旋转偏移随机扰动"].value()
            }
        }
    
    def on_view_changed(self):
        """视图切换时更新预览"""
        if self.radio_preview.isChecked():
            self.stack.setCurrentIndex(1)
            self.update_preview()
        else:
            self.stack.setCurrentIndex(0)

    def resizeEvent(self, event):
        """窗口大小改变时更新预览"""
        super().resizeEvent(event)
        if self.radio_preview.isChecked() and hasattr(self, 'preview'):
            # 在窗口调整大小后略微延迟更新，以确保UI已重新布局
            QTimer.singleShot(100, lambda: self.preview.fit_to_window())

    def clear_text(self):
        """清空文本内容"""
        reply = QMessageBox.question(self, "确认清空", 
                                   "确定要清空所有文本吗？这个操作不能撤销。",
                                   QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            self.text_edit.clear()
            self.statusBar().showMessage("文本已清空")
    
    def toggle_theme(self):
        """切换主题"""
        self.theme_manager.toggle_theme()
        theme_name = "深色" if self.theme_manager.current_theme == "dark" else "浅色"
        self.statusBar().showMessage(f"已切换到{theme_name}主题")
    
    def load_preset_dialog(self):
        """打开加载预设对话框"""
        presets = self.settings_manager.load_presets()
        if not presets:
            QMessageBox.information(self, "提示", "没有可用的预设。请先保存一个预设。")
            return
            
        preset_name, ok = QInputDialog.getItem(
            self, "加载预设", "选择预设:", 
            list(presets.keys()), 0, False
        )
        
        if ok and preset_name:
            self.load_preset(preset_name)
    
    def save_preset_dialog(self):
        """打开保存预设对话框"""
        name, ok = QInputDialog.getText(self, "保存预设", "请输入预设名称:")
        if ok and name:
            settings = self.get_current_settings()
            if self.settings_manager.save_preset(name, settings):
                QMessageBox.information(self, "成功", f"预设 '{name}' 已保存。")
            else:
                QMessageBox.warning(self, "错误", "保存预设失败。")
    
    def load_preset(self, preset_name):
        """加载指定的预设"""
        presets = self.settings_manager.load_presets()
        if preset_name not in presets:
            QMessageBox.warning(self, "错误", f"找不到预设 '{preset_name}'。")
            return
            
        preset = presets[preset_name]
        
        try:
            # 设置字体
            if 'font' in preset:
                index = self.font_combo.findText(preset['font'])
                if index >= 0:
                    self.font_combo.setCurrentIndex(index)
            
            # 设置字体大小
            if 'font_size' in preset:
                self.distortion_inputs["字体大小"].setValue(preset['font_size'])
            
            # 设置边距
            if 'margins' in preset:
                for name, value in preset['margins'].items():
                    if name in self.margin_inputs:
                        self.margin_inputs[name].setValue(value)
            
            # 设置扰动
            if 'distortions' in preset:
                for name, value in preset['distortions'].items():
                    mapped_name = {
                        "字间距扰动": "字间距扰动",
                        "行间距扰动": "行间距扰动",
                        "字体大小扰动": "字体大小随机扰动",
                        "横向偏移扰动": "笔画横向偏移随机扰动",
                        "纵向偏移扰动": "笔画纵向偏移随机扰动",
                        "旋转角度扰动": "笔画旋转偏移随机扰动"
                    }.get(name, name)
                    
                    if mapped_name in self.distortion_inputs:
                        self.distortion_inputs[mapped_name].setValue(value)
            
            self.statusBar().showMessage(f"已加载预设 '{preset_name}'")
            # 更新预览
            self.delayed_preview_update()
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"加载预设时出错: {str(e)}")
    
    def show_help(self):
        """显示帮助信息"""
        help_text = """
        <h2>手写模拟器使用帮助</h2>
        
        <h3>基本操作</h3>
        <ul>
            <li><b>导入文件</b>: 支持TXT、DOCX和PDF格式</li>
            <li><b>导出</b>: 支持PNG、JPEG、PDF和DOCX格式</li>
            <li><b>保存设置</b>: 保存当前设置供下次使用</li>
            <li><b>预设管理</b>: 保存和加载不同风格</li>
        </ul>
        
        <h3>参数调整</h3>
        <ul>
            <li><b>边距设置</b>: 调整文本上下左右边距和行间距、字间距</li>
            <li><b>扰动设置</b>: 调整文字随机性，使手写效果更自然</li>
            <li><b>字体大小</b>: 调整文字大小</li>
        </ul>
        
        <h3>批量处理</h3>
        <ul>
            <li><b>批量导入</b>: 一次导入多个文件并合并</li>
            <li><b>批量导出</b>: 导出多种格式或将长文本分块导出</li>
        </ul>
        
        <h3>小技巧</h3>
        <ul>
            <li>调整扰动参数可以控制手写风格随机程度</li>
            <li>选择不同背景可以模拟不同类型纸张</li>
            <li>预览功能可以快速查看效果，不必等待导出</li>
        </ul>
        """
        
        help_dialog = QDialog(self)
        help_dialog.setWindowTitle("使用帮助")
        help_dialog.resize(600, 400)
        
        layout = QVBoxLayout(help_dialog)
        text_browser = QTextBrowser()
        text_browser.setHtml(help_text)
        layout.addWidget(text_browser)
        
        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(help_dialog.accept)
        layout.addWidget(close_btn)
        
        help_dialog.exec()
    
    def show_about(self):
        """显示关于信息"""
        about_text = f"""
        <h2>手写模拟器 v1.0</h2>
        <p>这是一个基于Handright库的手写效果模拟工具，可以将电子文本转换为逼真的手写效果。</p>
        
        <p><b>主要功能：</b></p>
        <ul>
            <li>支持多种文件格式的导入和导出</li>
            <li>高度可定制的手写效果参数</li>
            <li>预设管理系统，方便保存和加载不同风格</li>
            <li>支持批量处理和多线程渲染</li>
        </ul>
        
        <p><b>技术支持：</b> 赵伟老师</p>
        <p><b>当前版本：</b> 1.0</p>
        <p><b>最后更新：</b> {time.strftime('%Y-%m-%d')}</p>
        """
        
        QMessageBox.about(self, "关于手写模拟器", about_text)
    
    def batch_import(self):
        """批量导入文件"""
        file_names, _ = QFileDialog.getOpenFileNames(
            self,
            "批量导入文件",
            "",
            FileManager.get_supported_import_formats()
        )
        
        if not file_names:
            return
            
        # 创建进度对话框
        progress = QProgressDialog("正在导入文件...", "取消", 0, len(file_names), self)
        progress.setWindowModality(Qt.WindowModal)
        progress.show()
        
        combined_text = ""
        imported_count = 0
        
        for i, file_name in enumerate(file_names):
            try:
                text = FileManager.import_file(Path(file_name))
                # 如果不是第一个文件，添加分隔符
                if combined_text:
                    combined_text += "\n\n" + "=" * 30 + "\n\n"
                combined_text += text
                imported_count += 1
                
                # 更新进度
                progress.setValue(i + 1)
                
                # 检查是否取消
                if progress.wasCanceled():
                    break
                    
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "导入错误",
                    f"导入文件 '{Path(file_name).name}' 时出错: {str(e)}"
                )
        
        # 更新文本框
        if imported_count > 0:
            # 如果已经有文本，询问是否替换或追加
            if self.text_edit.toPlainText():
                reply = QMessageBox.question(
                    self,
                    "导入方式",
                    "已有文本，您希望如何处理？",
                    QMessageBox.Cancel | QMessageBox.Discard | QMessageBox.Save,
                    QMessageBox.Save
                )
                
                if reply == QMessageBox.Save:  # 追加
                    self.text_edit.append("\n\n" + combined_text)
                elif reply == QMessageBox.Discard:  # 替换
                    self.text_edit.setPlainText(combined_text)
                # Cancel 则不做任何操作
            else:
                self.text_edit.setPlainText(combined_text)
                
            self.statusBar().showMessage(f"已导入 {imported_count} 个文件")
    
    def batch_export(self):
        """批量导出多种格式"""
        text = self.text_edit.toPlainText()
        if not text:
            QMessageBox.warning(self, "警告", "请先输入要转换的文本")
            return
        
        # 创建多格式导出对话框
        export_dialog = QDialog(self)
        export_dialog.setWindowTitle("批量导出")
        export_dialog.resize(400, 300)
        
        layout = QVBoxLayout(export_dialog)
        
        # 导出格式选择
        formats_group = QGroupBox("选择导出格式")
        formats_layout = QVBoxLayout(formats_group)
        
        pdf_check = QCheckBox("PDF")
        png_check = QCheckBox("PNG")
        jpg_check = QCheckBox("JPEG")
        docx_check = QCheckBox("DOCX")
        
        pdf_check.setChecked(True)
        
        formats_layout.addWidget(pdf_check)
        formats_layout.addWidget(png_check)
        formats_layout.addWidget(jpg_check)
        formats_layout.addWidget(docx_check)
        
        layout.addWidget(formats_group)
        
        # 导出选项
        options_group = QGroupBox("导出选项")
        options_layout = QFormLayout(options_group)
        
        dpi_spin = QSpinBox()
        dpi_spin.setRange(72, 600)
        dpi_spin.setValue(300)
        dpi_spin.setSingleStep(10)
        
        split_check = QCheckBox("拆分长文本")
        split_spin = QSpinBox()
        split_spin.setRange(100, 10000)
        split_spin.setValue(1000)
        split_spin.setSingleStep(100)
        split_spin.setSuffix(" 字/段")
        split_spin.setEnabled(False)
        
        split_check.toggled.connect(split_spin.setEnabled)
        
        options_layout.addRow("PDF分辨率:", dpi_spin)
        options_layout.addRow(split_check)
        options_layout.addRow("每段字数:", split_spin)
        
        layout.addWidget(options_group)
        
        # 按钮
        buttons = QHBoxLayout()
        export_btn = QPushButton("导出")
        cancel_btn = QPushButton("取消")
        
        export_btn.clicked.connect(export_dialog.accept)
        cancel_btn.clicked.connect(export_dialog.reject)
        
        buttons.addWidget(export_btn)
        buttons.addWidget(cancel_btn)
        
        layout.addLayout(buttons)
        
        # 显示对话框
        if export_dialog.exec() != QDialog.Accepted:
            return
        
        # 检查选择的格式
        formats = []
        if pdf_check.isChecked():
            formats.append("pdf")
        if png_check.isChecked():
            formats.append("png")
        if jpg_check.isChecked():
            formats.append("jpeg")
        if docx_check.isChecked():
            formats.append("docx")
        
        if not formats:
            QMessageBox.warning(self, "警告", "请至少选择一种导出格式")
            return
        
        # 获取导出目录
        export_dir = QFileDialog.getExistingDirectory(
            self,
            "选择导出目录",
            "",
            QFileDialog.ShowDirsOnly
        )
        
        if not export_dir:
            return
            
        # 准备导出
        try:
            # 获取当前设置并检查设置是否合理
            settings = self.get_current_settings()
            
            # 检查字体大小和行间距
            font_size = settings['font_size']
            line_spacing = settings['margins']['行间距']
            
            if line_spacing <= font_size:
                # 自动调整行间距
                old_line_spacing = line_spacing
                line_spacing = int(font_size * 1.5)
                settings['margins']['行间距'] = line_spacing
                
                # 更新UI中的行间距值
                if '行间距' in self.margin_inputs:
                    self.margin_inputs['行间距'].setValue(line_spacing)
                
                # 显示警告
                QMessageBox.warning(
                    self,
                    "参数已调整",
                    f"行间距值({old_line_spacing})小于字体大小({font_size})，已自动调整为{line_spacing}。"
                )
            
            # 创建模板
            template = self.create_template(settings)
            
            # 创建进度对话框
            total_tasks = len(formats)
            if split_check.isChecked():
                # 如果需要拆分，计算段数
                chars_per_section = split_spin.value()
                sections = (len(text) + chars_per_section - 1) // chars_per_section
                total_tasks *= sections
            
            progress = QProgressDialog("正在生成导出文件...", "取消", 0, total_tasks, self)
            progress.setWindowModality(Qt.WindowModal)
            progress.show()
            
            # 显示进度条
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            
            all_files = []
            current_task = 0
            
            # 处理文本拆分
            if split_check.isChecked():
                chars_per_section = split_spin.value()
                text_sections = []
                
                # 拆分文本
                for i in range(0, len(text), chars_per_section):
                    text_sections.append(text[i:i+chars_per_section])
            else:
                text_sections = [text]
            
            # 创建错误日志
            errors = []
            
            # 导出每种格式
            for section_idx, section_text in enumerate(text_sections):
                for fmt in formats:
                    # 更新进度
                    progress.setValue(current_task)
                    progress.setLabelText(f"正在导出 {fmt.upper()} 格式 (段落 {section_idx + 1}/{len(text_sections)})...")
                    
                    # 检查是否取消
                    if progress.wasCanceled():
                        break
                        
                    try:
                        # 生成手写图片
                        images = list(handwrite(section_text, template))
                        
                        if not images:
                            errors.append(f"段落 {section_idx + 1} 的 {fmt.upper()} 格式导出失败: 未能生成图像")
                            continue
                        
                        # 创建导出文件名
                        filename = f"手写_{time.strftime('%Y%m%d_%H%M%S')}"
                        if len(text_sections) > 1:
                            filename += f"_段落{section_idx + 1}"
                        
                        # 导出文件
                        base_path = os.path.join(export_dir, filename)
                        exported = FileManager.export_images(
                            images, 
                            base_path, 
                            format=fmt,
                            dpi=dpi_spin.value()
                        )
                        all_files.extend(exported)
                    
                    except Exception as e:
                        errors.append(f"段落 {section_idx + 1} 的 {fmt.upper()} 格式导出失败: {str(e)}")
                        
                    finally:
                        current_task += 1
                        progress.setValue(current_task)
            
            # 隐藏进度条
            self.progress_bar.setVisible(False)
            
            # 更新进度条完成
            progress.setValue(total_tasks)
            
            # 显示成功消息和错误信息
            result_dialog = QDialog(self)
            result_dialog.setWindowTitle("导出结果")
            result_dialog.resize(600, 400)
            
            result_layout = QVBoxLayout(result_dialog)
            
            # 导出结果信息
            header_text = f"共导出 {len(all_files)} 个文件"
            if errors:
                header_text += f"，但有 {len(errors)} 个错误"
            
            label = QLabel(header_text)
            result_layout.addWidget(label)
            
            # 创建选项卡控件
            tabs = QTabWidget()
            
            # 文件列表选项卡
            files_tab = QWidget()
            files_layout = QVBoxLayout(files_tab)
            file_list = QListWidget()
            for f in all_files:
                file_list.addItem(f)
            files_layout.addWidget(file_list)
            
            # 错误列表选项卡
            errors_tab = QWidget()
            errors_layout = QVBoxLayout(errors_tab)
            error_list = QListWidget()
            for e in errors:
                error_list.addItem(e)
            errors_layout.addWidget(error_list)
            
            # 添加选项卡
            tabs.addTab(files_tab, f"文件列表 ({len(all_files)})")
            if errors:
                tabs.addTab(errors_tab, f"错误信息 ({len(errors)})")
            
            result_layout.addWidget(tabs)
            
            # 底部按钮
            buttons = QHBoxLayout()
            
            # 打开目录按钮
            open_dir_btn = QPushButton("打开导出目录")
            open_dir_btn.clicked.connect(lambda: os.startfile(export_dir))
            
            close_btn = QPushButton("关闭")
            close_btn.clicked.connect(result_dialog.accept)
            
            buttons.addWidget(open_dir_btn)
            buttons.addWidget(close_btn)
            result_layout.addLayout(buttons)
            
            result_dialog.exec()
            
        except Exception as e:
            self.progress_bar.setVisible(False)
            QMessageBox.critical(self, "导出错误", str(e))

def main():
    try:
        app = QApplication(sys.argv)
        
        # 设置应用样式
        app.setStyle("Fusion")
        
        # 应用内置的默认样式
        StyleManager.apply_default_style(app)
        
        # 设置应用名称和组织信息
        app.setApplicationName("手写模拟器")
        app.setOrganizationName("HandwriteSim")
        
        # 检查必要的目录是否存在
        for dir_name in ["fonts", "backgrounds"]:
            dir_path = Path(dir_name)
            if not dir_path.exists():
                dir_path.mkdir(exist_ok=True)
                print(f"已创建目录: {dir_name}")
        
        # 检查字体目录是否有字体文件
        if not list(Path("fonts").glob("*.ttf")):
            print("提示: 字体目录中没有找到TTF字体文件，将使用系统默认字体")
            
            # 尝试使用默认的示例字体
            if Path("Handright-master.zip").exists():
                try:
                    import zipfile
                    with zipfile.ZipFile("Handright-master.zip", 'r') as zip_ref:
                        for file in zip_ref.namelist():
                            if file.endswith('.ttf') and not file.startswith('__MACOSX'):
                                # 提取到fonts目录
                                zip_ref.extract(file, "temp")
                                # 移动文件
                                import shutil
                                source_file = Path("temp") / file
                                target_file = Path("fonts") / Path(file).name
                                shutil.move(str(source_file), str(target_file))
                                print(f"已从示例包中提取字体: {target_file}")
                    # 清理临时目录
                    import shutil
                    if Path("temp").exists():
                        shutil.rmtree("temp")
                except Exception as e:
                    print(f"提取示例字体失败: {e}")
        
        # 创建主窗口并显示
        window = MainWindow()
        window.resize(900, 500)
        window.show()
        
        # 设置窗口图标
        try:
            window.setWindowIcon(QIcon("icon.png"))
        except:
            pass  # 图标不存在不影响程序运行
        
        # 提示版本信息
        print(f"手写模拟器 v1.0 已启动")
        print(f"Python版本: {sys.version}")
   
        
        # 执行应用程序
        return app.exec()
        
    except Exception as e:
        print(f"启动失败: {e}")
        # 如果图形界面已启动，显示错误对话框
        try:
            QMessageBox.critical(None, "启动错误", f"程序启动失败:\n{str(e)}")
        except:
            pass
        return 1

if __name__ == "__main__":
    sys.exit(main())
